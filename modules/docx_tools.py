import os
from docx import Document
from datetime import datetime

def add_comments_to_docx(input_path, issues, output_dir="outputs"):
    """
    Adds comments to a DOCX file based on compliance issues.

    Args:
        input_path (str): Path to the original .docx file.
        issues (list[dict]): List of issues with keys:
            - document
            - section (optional)
            - issue
            - severity
            - suggestion
        output_dir (str): Where to save reviewed file.

    Returns:
        str: Path to the reviewed .docx file.
    """
    if not os.path.exists(input_path):
        raise FileNotFoundError(f"File not found: {input_path}")

    os.makedirs(output_dir, exist_ok=True)
    doc = Document(input_path)

    # For each issue, try to find where to insert a comment
    for issue in issues:
        issue_text = f"[{issue.get('severity', 'Info')}] {issue.get('issue', '')} | Suggestion: {issue.get('suggestion', '')}"
        found = False

        # If section is given, search for it
        if issue.get("section"):
            for para in doc.paragraphs:
                if issue["section"].lower() in para.text.lower():
                    para.add_run(f"  <-- COMMENT: {issue_text}").italic = True
                    found = True
                    break
        
        # If no section match, try a keyword from the issue text
        if not found:
            for para in doc.paragraphs:
                if any(word.lower() in para.text.lower() for word in issue["issue"].split()[:3]):
                    para.add_run(f"  <-- COMMENT: {issue_text}").italic = True
                    found = True
                    break
        
        # If still not found, just append comment at the end of the doc
        if not found:
            doc.add_paragraph(f"COMMENT (Unplaced): {issue_text}")

    # Save reviewed file
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    reviewed_path = os.path.join(output_dir, f"reviewed_{timestamp}.docx")
    doc.save(reviewed_path)
    return reviewed_path

if __name__ == "__main__":
    # Example test
    sample_issues = [
        {
            "document": "Articles of Association",
            "section": "Clause 3.1",
            "issue": "Jurisdiction clause does not specify ADGM",
            "severity": "High",
            "suggestion": "Update jurisdiction to ADGM Courts."
        }
    ]
    reviewed = add_comments_to_docx("data/sample_docs/sample_aoa.docx", sample_issues)
    print("Reviewed file saved at:", reviewed)
